import os
from docx import Document
from fpdf import FPDF

DOCX_PATH = "docs/INFORME_PRUEBAS_UNITARIAS.docx"
PDF_PATH = "docs/INFORME_PRUEBAS_UNITARIAS.pdf"
ARIAL_PATH = "C:/Windows/Fonts/arial.ttf"
ARIAL_BOLD_PATH = "C:/Windows/Fonts/arialbd.ttf"
ARIAL_ITALIC_PATH = "C:/Windows/Fonts/ariali.ttf"
ARIAL_BI_PATH = "C:/Windows/Fonts/arialbi.ttf"


class DocxToPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font("Arial", "I", 8)
            self.cell(0, 5, "INFORME PRUEBAS UNITARIAS - Calletano POS", align="C")
            self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", align="C")


def get_heading_level(style_name):
    name = style_name.lower()
    if "title" in name:
        return 1
    if "heading" in name:
        for c in name:
            if c.isdigit():
                return int(c)
        return 1
    return 0


def replace_unsupported_chars(text):
    """Replace emojis and special characters with text equivalents."""
    replacements = {
        "\u2705": "[OK]",    # check mark
        "\u274c": "[X]",     # cross mark
        "\u2714": "[ok]",    # heavy check
        "\u2716": "[x]",     # heavy multiplication
        "\u2713": "[v]",     # check
        "\u2795": "[+]",     # plus
        "\u2796": "[-]",     # minus
        "\u2b50": "[*]",     # star
        "\u2139": "[i]",     # info
        "\u26a0": "[!]",     # warning
        "\u200b": "",        # zero-width space
        "\u200c": "",        # zero-width non-joiner
        "\u200d": "",        # zero-width joiner
        "\uf0b7": "-",       # bullet
        "\u2022": "-",       # bullet
        "\u2023": ">",       # triangular bullet
        "\u25e6": "o",       # white bullet
        "\u2043": "-",       # hyphen bullet
        "\u2212": "-",       # minus sign
        "\u2013": "-",       # en dash
        "\u2014": "--",      # em dash
        "\u2018": "'",       # left single quote
        "\u2019": "'",       # right single quote
        "\u201c": "\"",      # left double quote
        "\u201d": "\"",      # right double quote
        "\u00a0": " ",       # non-breaking space
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text


def convert():
    if not os.path.exists(DOCX_PATH):
        print(f"Error: No se encuentra el archivo {DOCX_PATH}")
        return False

    doc = Document(DOCX_PATH)
    pdf = DocxToPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Add Unicode-capable fonts from Windows
    pdf.add_font("Arial", "", ARIAL_PATH, uni=True)
    pdf.add_font("Arial", "B", ARIAL_BOLD_PATH, uni=True)
    pdf.add_font("Arial", "I", ARIAL_ITALIC_PATH, uni=True)
    pdf.add_font("Arial", "BI", ARIAL_BI_PATH, uni=True)

    pdf.add_page()

    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(4)
            continue

        # Replace characters not supported by the font
        text = replace_unsupported_chars(text)

        level = get_heading_level(para.style.name if para.style else "")

        if level == 1:
            pdf.set_font("Arial", "B", 16)
            pdf.multi_cell(page_w, 8, text)
            pdf.ln(3)
        elif level == 2:
            pdf.set_font("Arial", "B", 13)
            pdf.multi_cell(page_w, 7, text)
            pdf.ln(2)
        elif level == 3:
            pdf.set_font("Arial", "B", 12)
            pdf.multi_cell(page_w, 6, text)
            pdf.ln(1)
        else:
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            font_style = "B" if is_bold else ""
            pdf.set_font("Arial", font_style, 11)

            starts_with_bullet = text.startswith("- ") or text.startswith("* ")
            if not starts_with_bullet and len(text) > 0 and ord(text[0]) in [8226, 8211, 8227]:
                starts_with_bullet = True

            indent = 5 if starts_with_bullet else 0
            pdf.set_x(pdf.l_margin + indent)
            pdf.multi_cell(page_w - indent, 6, text)

    # Extract tables
    for table in doc.tables:
        pdf.ln(3)
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for cell in row.cells:
                cells_text.append(replace_unsupported_chars(cell.text.strip()))
            row_line = "  |  ".join(cells_text)
            font_style = "B" if row_idx == 0 else ""
            pdf.set_font("Arial", font_style, 10)
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(page_w - 6, 5, row_line)
            pdf.ln(1)

    pdf.output(PDF_PATH)
    print(f"PDF creado exitosamente: {PDF_PATH}")
    print(f"Tamano: {os.path.getsize(PDF_PATH)} bytes")
    return True


if __name__ == "__main__":
    convert()
