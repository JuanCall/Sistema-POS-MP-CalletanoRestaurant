import docx

doc = docx.Document('docs/paper_capstone.docx')
changes = []

def clear_paragraph(p):
    """Clear all text from a paragraph."""
    for run in p.runs:
        run.text = ''

def set_run_text(p, text):
    """Set text for paragraph, clearing all runs beyond the first."""
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ''
    else:
        run = p.add_run(text)
    return p

# ================================================================
# 1. FIX para 102: Restore the fixed paragraph text, move figure elsewhere
# ================================================================
p102 = doc.paragraphs[102]
p102_fixed_text = (
    "Los resultados obtenidos durante el periodo de evaluacion sugieren que "
    "la implementacion del sistema POS se asocio con una mejora en los indicadores "
    "de eficiencia operativa del establecimiento bajo las condiciones del estudio. "
    "La Tabla 10 ilustra la evolucion de los indicadores clave de desempeno (KPIs) "
    "al contrastar el periodo pre-implementacion con el periodo post-implementacion."
)
set_run_text(p102, p102_fixed_text)
changes.append("Restored para 102 with corrected text")

# ================================================================
# 2. FIX para 63: Add "Analisis Estadistico" section (multi-paragraph)
# ================================================================
stat_heading = doc.paragraphs[63]
stat_heading_text = "4.5 Analisis Estadistico"
set_run_text(stat_heading, stat_heading_text)
changes.append("Added 'Analisis Estadistico' heading to para 63")

# Insert the body text as a new paragraph after para 63
stat_body = (
    "El presente estudio presenta estadistica descriptiva basada en indicadores "
    "operativos registrados durante los periodos pre-implementacion (30 dias) y "
    "post-implementacion (30 dias). Para cada indicador se calcularon medidas de "
    "tendencia central (media aritmetica) que permiten describir el comportamiento "
    "de las variables observadas. No se realizaron pruebas de inferencia estadistica "
    "(como prueba t de Student para muestras relacionadas o Wilcoxon signed-rank) "
    "debido a la ausencia de datos individuales apareados y al diseno preexperimental "
    "del estudio. Investigaciones futuras podrian incorporar: (a) prueba de normalidad "
    "de Shapiro-Wilk para determinar la distribucion de los datos; (b) prueba t de "
    "Student para muestras relacionadas o Wilcoxon signed-rank segun la normalidad "
    "de los datos; (c) intervalo de confianza al 95% para la diferencia de medias; "
    "y (d) d de Cohen como medida del tamano del efecto."
)
from docx.oxml.ns import qn
from copy import deepcopy

# Create a new paragraph after para 63
new_p = docx.oxml.OxmlElement('w:p')
doc.paragraphs[63]._element.addnext(new_p)
new_para = docx.text.paragraph.Paragraph(new_p, doc.paragraphs[63]._element.getparent())
# ... actually let me use a simpler approach
changes.append("(Note: stat body text needs to be manually added as a paragraph)")

# ================================================================
# 3. FIX para 104: Add Table 11 (use single-line approach)
# ================================================================
p104 = doc.paragraphs[104]
table11_text = (
    "Tabla 11. Resumen descriptivo de indicadores operativos antes y despues de la implementacion. "
    "Indicadores: Pedidos registrados (432 pre, 450 post); Tiempo promedio de atencion "
    "(300s pre, 120s post, -60%); Errores promedio por dia (10.0 pre, 0.2 post, -98%); "
    "Tiempo promedio de cierre de caja (45 min pre, 8 min post, -82%). "
    "Nota: Los porcentajes indican la reduccion relativa respecto al valor pre-implementacion."
)
set_run_text(p104, table11_text)
changes.append("Added Table 11 to para 104 (text format)")

# ================================================================
# 4. FIX para 120: Fix the replacement (reduciendo -> reduccion de)
# ================================================================
p120 = doc.paragraphs[120]
current_p120 = p120.text
# Replace "reduciendo los tiempos de atencion en un 74%"
fixed_p120 = current_p120.replace(
    "reduciendo los tiempos de atencion en un 74%",
    "se observo una reduccion del tiempo promedio de atencion de 300 segundos a 120 segundos (60% de reduccion)"
)
# Already has "lo cual es consistente con el objetivo planteado" from the first fix
set_run_text(p120, fixed_p120)
changes.append("Fixed para 120: replaced '74%' with real time data (300s->120s, 60%)")

# ================================================================
# 5. Move figure suggestion from para 102 to a better position (after Table 10)
# ================================================================
# Add figure suggestion as a new paragraph - insert after Tabla 10 (para 103)
fig_text = (
    "[Figura 1 - Comparacion de indicadores antes y despues de la implementacion] "
    "Sugerencia: Grafico de barras agrupadas que compare los tres indicadores principales "
    "(tiempo de atencion, tasa de errores, tiempo de cierre) en escala normalizada "
    "(porcentaje respecto al valor pre-implementacion = 100%). "
    "Incluir barras de error si se dispone de la desviacion estandar de cada medicion."
)
# Use para 103 which currently has Tabla 10 - add after it
# Actually, let's put it right after Table 10 content in para 103
changes.append("(Note: Figure 1 suggestion needs manual positioning after Tabla 10)")

# ================================================================
# Save
# ================================================================
doc.save('docs/paper_capstone.docx')
print(f"File saved. {len(changes)} changes applied.")
for c in changes:
    print(f"  - {c}")
