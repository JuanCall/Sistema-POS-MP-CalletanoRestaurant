import docx
from docx.oxml.ns import qn

doc = docx.Document('docs/paper_capstone.docx')
changes = []

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)

def find_para(substr):
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i, p
    return None, None

# ================================================================
# FIX 1: Replace "74%" with real time data in conclusions
# ================================================================
_, p = find_para("reduciendo los tiempos de aten")
if p and "74%" in p.text:
    set_text(p, p.text.replace(
        "reduciendo los tiempos de atenció\u0301n en un 74%",
        "se observo una reduccion del tiempo promedio de atencion de 300 segundos a 120 segundos (60% de reduccion)"
    ).replace(
        "reduciendo los tiempos de atenci\u00f3n en un 74%",
        "se observo una reduccion del tiempo promedio de atencion de 300 segundos a 120 segundos (60% de reduccion)"
    ).replace(
        "74%", "60%"
    ))
    changes.append("FIX 1: Replaced '74%' with '60% de reduccion' in conclusions")

# ================================================================
# FIX 2: Fix cash close times (60->10 -> 45->8)
# ================================================================
_, p = find_para("reduciendo el tiempo administrativo diario")
if p and "45 a 8 minutos" not in p.text:
    set_text(p, p.text.replace(
        "reduciendo el tiempo administrativo diario de 60 a 10 minutos",
        "reduciendo el tiempo administrativo diario de 45 a 8 minutos promedio"
    ).replace(
        "de 60 a 10 minutos",
        "de 45 a 8 minutos"
    ))
    changes.append("FIX 2: Updated cash close times to 45->8 min")

# ================================================================
# FIX 3: Clean up Section 9 structure
# ================================================================
# Find current positions
_, p_lim_head = find_para("9.2 Limitaciones del Estudio")
_, p_amen_head = find_para("9.3 Amenazas a la Validez")
_, p_trab_content = find_para("Las lineas de investigacion")

# Clear duplicate headings and stray recomendaciones
for _, p in [find_para("Actualizaci")]:
    if p:
        set_text(p, "")
        changes.append("FIX 3: Cleared stray 'Actualizacion'")

_, p_esc = find_para("Escalabilidad: Integrar")
if p_esc:
    set_text(p_esc, "")
    changes.append("FIX 3: Cleared stray 'Escalabilidad'")

_, p_cap = find_para("Capacitaci\u00f3n")
if p_cap:
    set_text(p_cap, "")
    changes.append("FIX 3: Cleared stray 'Capacitacion'")

# Check if Trabajo Futuro heading exists
_, p_tf_head = find_para("9.4 Trabajo Futuro")
if not p_tf_head:
    _, p_tf_head = find_para("Trabajo Futuro")

if not p_tf_head and p_trab_content:
    # Find the paragraph just before the trabajo futuro content
    i_tf_content, _ = find_para("Las lineas de investigacion")
    if i_tf_content:
        # Find the Amenazas paragraph to insert after it
        i_amen, p_amen = find_para("9.3 Amenazas")
        if i_amen is not None:
            new_p = docx.oxml.OxmlElement('w:p')
            new_r = docx.oxml.OxmlElement('w:r')
            new_t = docx.oxml.OxmlElement('w:t')
            new_t.text = "9.4 Trabajo Futuro"
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            new_p.append(new_r)
            doc.paragraphs[i_amen]._element.addnext(new_p)
            changes.append("FIX 3: Added '9.4 Trabajo Futuro' heading")

# ================================================================
# Save
# ================================================================
doc.save('docs/paper_capstone.docx')
print(f"OK - {len(changes)} changes")
for c in changes:
    print(f"  [x] {c}")
