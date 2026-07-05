import docx
from docx.oxml.ns import qn

doc = docx.Document('docs/paper_capstone.docx')
changes = []

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)

def insert_para_after(ref_para, text):
    """Insert a new paragraph after the reference paragraph."""
    new_p_elm = docx.oxml.OxmlElement('w:p')
    new_r = docx.oxml.OxmlElement('w:r')
    new_t = docx.oxml.OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p_elm.append(new_r)
    ref_para._element.addnext(new_p_elm)
    return docx.text.paragraph.Paragraph(new_p_elm, ref_para._element.getparent())

# ================================================================
# 1. FIX para 102: Restore corrected text (figure overwrote it)
# ================================================================
set_text(doc.paragraphs[102],
    "Los resultados obtenidos durante el periodo de evaluacion sugieren que "
    "la implementacion del sistema POS se asocio con una mejora en los indicadores "
    "de eficiencia operativa del establecimiento bajo las condiciones del estudio. "
    "La Tabla 10 ilustra la evolucion de los indicadores clave de desempeno (KPIs) "
    "al contrastar el periodo pre-implementacion con el periodo post-implementacion."
)
changes.append("Fixed para 102: restored corrected text")

# ================================================================
# 2. FIX para 63 heading + add body paragraph after 63
# ================================================================
set_text(doc.paragraphs[63], "4.5 Analisis Estadistico")
changes.append("Added heading '4.5 Analisis Estadistico' to para 63")

body = (
    "El presente estudio presenta estadistica descriptiva basada en indicadores "
    "operativos registrados durante los periodos pre-implementacion (30 dias) y "
    "post-implementacion (30 dias). Para cada indicador se calcularon medidas de "
    "tendencia central (media aritmetica) que permiten describir el comportamiento "
    "de las variables observadas. No se realizaron pruebas de inferencia estadistica "
    "(como prueba t de Student para muestras relacionadas o Wilcoxon signed-rank) "
    "debido a la ausencia de datos individuales apareados y al diseno preexperimental "
    "del estudio. Investigaciones futuras podrian incorporar: (a) prueba de normalidad "
    "de Shapiro-Wilk para determinar la distribucion de los datos; (b) prueba t de "
    "Student para muestras relacionadas o Wilcoxon signed-rank segun la normalidad "
    "de los datos; (c) intervalo de confianza al 95% para la diferencia de medias; "
    "y (d) d de Cohen como medida del tamano del efecto."
)
insert_para_after(doc.paragraphs[63], body)
changes.append("Added 'Analisis Estadistico' body paragraph after para 63")

# ================================================================
# 3. FIX para 104: Add Table 11 (empty, should have resumen table)
# ================================================================
set_text(doc.paragraphs[104],
    "Tabla 11. Resumen descriptivo de indicadores operativos antes y despues "
    "de la implementacion. Indicadores: Pedidos registrados (432 pre, 450 post); "
    "Tiempo promedio de atencion (300s pre, 120s post, -60%); Errores promedio "
    "por dia (10.0 pre, 0.2 post, -98%); Tiempo promedio de cierre de caja "
    "(45 min pre, 8 min post, -82%). Nota: Los porcentajes indican la reduccion "
    "relativa respecto al valor pre-implementacion."
)
changes.append("Added Table 11 to para 104")

# ================================================================
# 4. FIX para 120: Fix replacement ("reduciendo" -> actual data)
# ================================================================
current = doc.paragraphs[120].text
fixed = current.replace(
    "reduciendo los tiempos de atencion en un 74%",
    "se observo una reduccion del tiempo promedio de atencion de 300 segundos a 120 segundos (60% de reduccion)"
)
set_text(doc.paragraphs[120], fixed)
changes.append("Fixed para 120: replaced '74%' with real time data")

# ================================================================
# 5. Add Figure 1 suggestion after Tabla 10 (para 103) 
# ================================================================
fig_text = (
    "[Figura 1 - Comparacion de indicadores antes y despues de la implementacion] "
    "Sugerencia: Grafico de barras agrupadas que compare los tres indicadores principales "
    "(tiempo de atencion, tasa de errores, tiempo de cierre) en escala normalizada "
    "(porcentaje respecto al valor pre-implementacion = 100%). Incluir barras de error "
    "si se dispone de la desviacion estandar de cada medicion."
)
insert_para_after(doc.paragraphs[103], fig_text)
changes.append("Added Figure 1 suggestion after Tabla 10 (para 103)")

# ================================================================
# Save
# ================================================================
doc.save('docs/paper_capstone.docx')
print(f"OK - {len(changes)} changes applied")
for c in changes:
    print(f"  [x] {c}")
