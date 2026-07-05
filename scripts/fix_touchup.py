import docx
from docx.oxml.ns import qn

doc = docx.Document('docs/paper_capstone.docx')
changes = []

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''

# Find duplicates: first "9.2 Limitaciones" at [127] - clear it
count = 0
for i, p in enumerate(doc.paragraphs):
    if "9.2 Limitaciones del Estudio" in p.text:
        count += 1
        if count == 1:  # First occurrence = duplicate
            set_text(p, "")
            changes.append(f"Cleared duplicate '9.2 Limitaciones' at para [{i}]")
        else:
            changes.append(f"Kept second '9.2 Limitaciones' at para [{i}]")

# Find if "9.4 Trabajo Futuro" heading exists
has_heading = any("9.4 Trabajo Futuro" in p.text for p in doc.paragraphs)

if not has_heading:
    # Find the Amenazas paragraph to insert after it
    amen_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "9.3 Amenazas" in p.text:
            amen_idx = i
            break
    
    if amen_idx is not None:
        new_p = docx.oxml.OxmlElement('w:p')
        new_r = docx.oxml.OxmlElement('w:r')
        new_t = docx.oxml.OxmlElement('w:t')
        new_t.text = "9.4 Trabajo Futuro"
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        doc.paragraphs[amen_idx]._element.addnext(new_p)
        changes.append(f"Added '9.4 Trabajo Futuro' heading after para [{amen_idx}]")
    else:
        changes.append("WARN: Could not find '9.3 Amenazas' paragraph")

doc.save('docs/paper_capstone.docx')
print(f"OK - {len(changes)} changes")
for c in changes:
    print(f"  [x] {c}")
