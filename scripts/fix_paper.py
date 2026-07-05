import docx

doc = docx.Document('docs/paper_capstone.docx')
changes_made = []

# 1. REMOVE university course formatting (paragraph 4)
p = doc.paragraphs[4]
for run in p.runs:
    run.text = ''
changes_made.append("Removed: Docente del Curso / Capstone Project line")

# 2. REWRITE English Abstract (paragraph 12) - IMRaD structure
p12 = doc.paragraphs[12]
new_abstract = (
    "This study presents the design, development, and evaluation of a multiplatform Point-of-Sale (POS) system "
    "implemented at CALLETANO RESTAURANT in Mancora, Peru. The system addresses deficiencies in manual operational management, "
    "including delays in customer service, recurrent order registration errors, and limited financial control. "
    "A pre-experimental design (G \u2022 O1 \u2022 X \u2022 O2) was employed, with measurements taken before and after the intervention. "
    "The technological solution is based on a hybrid client-server architecture using Node.js and Express for the local backend, "
    "React Native for the mobile ordering application, Electron.js for the desktop cashier interface, SQLite for local persistence, "
    "and Firebase Firestore for cloud synchronization. Results obtained during the evaluation period indicate a 74% reduction "
    "in customer service time, elimination of 93% of order errors, and a decrease in daily cash closing time from 60 to 10 minutes. "
    "These findings support the hypothesis that the implementation of a multiplatform POS system positively influences "
    "the operational efficiency of the restaurant establishment under the conditions of this study."
)
p12.runs[0].text = new_abstract
for run in p12.runs[1:]:
    run.text = ''
changes_made.append("Rewrote English Abstract with IMRaD structure (Background, Objective, Methods, Results, Conclusion)")

# 3. FIX reference [2] in bibliography (paragraph 138)
p138 = doc.paragraphs[138]
p138.runs[0].text = "[2] Teerasoponpong, S., & Sopadang, A. (2022). Decision support system for adaptive sourcing and inventory management in small- and medium-sized enterprises. Robotics and Computer-Integrated Manufacturing, 73, 102226. https://doi.org/10.1016/j.rcim.2021.102226"
for run in p138.runs[1:]:
    run.text = ''
changes_made.append("Fixed reference [2]: corrected authors, journal title (English), article ID, and DOI")

# 4. FIX reference [3] in bibliography (paragraph 139)
p139 = doc.paragraphs[139]
p139.runs[0].text = "[3] Ghazi, M., Salih, H. S., & Aljanabi, M. (2023). Implementing an automated inventory management system for small and medium-sized enterprises. Iraqi Journal for Computer Science and Mathematics, 4(2), 238-244. https://doi.org/10.52866/ijcsm.2023.02.02.021"
for run in p139.runs[1:]:
    run.text = ''
changes_made.append("Fixed reference [3]: corrected authors, English title, volume, pages, and DOI")

# 5. REPLACE reference [6] in bibliography (paragraph 142) - FICTITIOUS
p142 = doc.paragraphs[142]
p142.runs[0].text = "[6] Delgado, A., Huaman\u00ed, E. L., & Diego, S. (2020). Design of web systems for inventory control in the E-commerce sector under the Agile methodologies approach. International Journal of Emerging Trends in Engineering Research, 8(7), 3129-3133. https://doi.org/10.30534/ijeter/2020/41872020"
for run in p142.runs[1:]:
    run.text = ''
changes_made.append("Replaced fictitious reference [6] with verified real paper on Agile + inventory systems (Delgado et al., 2020)")

# 6. REPLACE reference [7] in bibliography (paragraph 143) - FICTITIOUS
p143 = doc.paragraphs[143]
p143.runs[0].text = "[7] Jim\u00e9nez, L. V., Espinoza, H. G., Jim\u00e9nez, S. R., & Gamarra, J. (2024). Software design aimed at proper order management in SMEs. International Journal of Advanced Computer Science and Applications, 15(12), 86-94. https://doi.org/10.14569/IJACSA.2024.0151209"
for run in p143.runs[1:]:
    run.text = ''
changes_made.append("Replaced fictitious reference [7] with verified real paper on software design for SMEs (Jim\u00e9nez et al., 2024)")

# 7. FIX reference [8] format in bibliography (paragraph 144)
p144 = doc.paragraphs[144]
p144.runs[0].text = "[8] Mishra, A., & Alzoubi, Y. I. (2023). Structured software development versus agile software development: a comparative analysis. International Journal of System Assurance Engineering and Management, 14(4), 1504-1522. https://doi.org/10.1007/s13198-023-01958-5"
for run in p144.runs[1:]:
    run.text = ''
changes_made.append("Fixed reference [8]: corrected title to English original, proper author format, and DOI")

# 8. FIX citation in text paragraph [48] - reference [2]
p48 = doc.paragraphs[48]
old48 = p48.text
new48 = old48.replace(
    "Siravat y Apichat (2022) [2]",
    "Teerasoponpong y Sopadang (2022) [2]"
)
new48 = new48.replace(
    "sistemas de soporte a la toma de decisiones para la gesti\u00f3n adaptativa de aprovisionamientos e inventarios en PYMES, concluyendo",
    "un sistema de soporte a la toma de decisiones para la gesti\u00f3n adaptativa de aprovisionamientos e inventarios en PYMES. Sus resultados indicaron"
)
p48.runs[0].text = new48
for run in p48.runs[1:]:
    run.text = ''
changes_made.append("Fixed citation [2] in text: corrected author names and improved scientific language")

# 9. FIX citation in text paragraph [50] - reference [6] (fictitious replacement)
p50 = doc.paragraphs[50]
old50 = p50.text
new50 = old50.replace(
    "Contreras y Mu\u00f1oz (2019) [6] implementaron un sistema de gesti\u00f3n de inventarios en una tienda de retail utilizando metodolog\u00eda \u00e1gil, validando que el enfoque iterativo es superior al modelo en cascada para entornos comerciales con requisitos cambiantes.",
    "Delgado, Huaman\u00ed y Diego (2020) [6] dise\u00f1aron sistemas web para el control de inventarios en el sector comercio electr\u00f3nico utilizando un enfoque de metodolog\u00edas \u00e1giles, demostrando la aplicabilidad de Scrum en entornos de gesti\u00f3n de inventarios con requisitos cambiantes."
)
p50.runs[0].text = new50
for run in p50.runs[1:]:
    run.text = ''
changes_made.append("Replaced fictitious citation [6] in text with verified real paper content")

# 10. FIX absolute language in paragraph [106]
p106 = doc.paragraphs[106]
old106 = p106.text
new106 = old106.replace(
    "Los resultados demuestran que la transici\u00f3n de comandas f\u00edsicas en papel a la aplicaci\u00f3n React Native redujo el tiempo promedio de env\u00edo de pedidos a cocina de 4.2 minutos a menos de 5 segundos.",
    "Los resultados obtenidos durante el per\u00edodo evaluado indican que la transici\u00f3n de comandas f\u00edsicas en papel a la aplicaci\u00f3n React Native se asoci\u00f3 con una reducci\u00f3n del tiempo promedio de env\u00edo de pedidos a cocina, de 4.2 minutos a menos de 5 segundos."
)
if "eliminando las discrepancias" in new106:
    new106 = new106.replace(
        "eliminando las discrepancias de dinero que se daban en el cuadre de caja anal\u00f3gico.",
        "reduciendo las discrepancias de dinero que se presentaban en el cuadre de caja manual."
    )
p106.runs[0].text = new106
for run in p106.runs[1:]:
    run.text = ''
changes_made.append("Fixed absolute language in paragraph [106]: replaced affirmative claims with evidence-based phrasing")

# 11. FIX absolute language in paragraph [107]
p107 = doc.paragraphs[107]
old107 = p107.text
new107 = old107.replace(
    "La persistencia local de SQLite resolvi\u00f3 el problema latente en la zona de Mancora respecto a la inestabilidad de las redes m\u00f3viles e inal\u00e1mbricas de internet, garantizando cero p\u00e9rdida de transacciones durante cortes del ISP.",
    "La persistencia local con SQLite contribuy\u00f3 a mitigar las dificultades asociadas a la inestabilidad de las redes m\u00f3viles e inal\u00e1mbricas en la zona de Mancora. Durante el per\u00edodo de evaluaci\u00f3n no se registraron p\u00e9rdidas de transacciones durante los cortes del proveedor de servicios de internet."
)
new107 = new107.replace(
    "Los errores en pedidos y facturaci\u00f3n bajaron a 0%",
    "Los errores en pedidos y facturaci\u00f3n se redujeron a 0% durante el per\u00edodo de observaci\u00f3n"
)
new107 = new107.replace(
    "ya que los platos agotados se bloquean en tiempo real en las tablets de los mozos, impidiendo la venta de stock inexistente.",
    "ya que el sistema bloquea en tiempo real la selecci\u00f3n de platos agotados en las tablets de los mozos, previniendo la venta de stock no disponible."
)
p107.runs[0].text = new107
for run in p107.runs[1:]:
    run.text = ''
changes_made.append("Fixed absolute language in paragraph [107]: replaced guarantees with observed outcomes")

# 12. FIX absolute language in paragraph [115]
p115 = doc.paragraphs[115]
old115 = p115.text
new115 = old115.replace(
    "El proyecto es 100% viable.",
    "Los resultados del an\u00e1lisis de viabilidad indican que el proyecto es financieramente viable bajo las condiciones evaluadas."
)
new115 = new115.replace(
    "cero ca\u00eddas y cero latencia",
    "sin interrupciones significativas y con latencia m\u00ednima durante las pruebas realizadas"
)
new115 = new115.replace(
    "eliminando la dependencia cr\u00edtica de un proveedor de internet (ISP)",
    "reduciendo la dependencia cr\u00edtica de un proveedor de internet (ISP) durante la operaci\u00f3n local"
)
p115.runs[0].text = new115
for run in p115.runs[1:]:
    run.text = ''
changes_made.append("Fixed absolute language in paragraph [115]: removed marketing terms (100%, cero), added evidence-based qualifiers")

# SAVE
doc.save('docs/paper_capstone.docx')

print("File saved: docs/paper_capstone.docx")
print()
print("CHANGES MADE (" + str(len(changes_made)) + "):")
for c in changes_made:
    print("  - " + c)
